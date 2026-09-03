# -*- coding: utf-8 -*-
"""Ver1.1.10.6 データ集約ドキュメント（Docx）生成。利用者（操作者）向け。
Ver1.1.9.5 生成スクリプトをベースに、整形 DSL テスト画面・スキャン／一括／上限警告を追記。
出力先: docs/Ver1.1.10.6/
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent / "Ver1.1.10.6"
VERSION = "1.1.10.6"
DOC_DATE = "2026年9月"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x0D, 0x5C, 0x63)
GRAY = RGBColor(0x44, 0x44, 0x44)
HINT_BG = "FFF3CD"
HEAD_BG = "1F3A5F"
ALT_BG = "F4F7FA"
CODE_BG = "F0F0F0"


def _set_run_font(run, name="游ゴシック", size=10.5, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_border(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "B0B8C1")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def new_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "游ゴシック"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
    pf = styles["Normal"].paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return doc


def add_footer(doc: Document, title: str) -> None:
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{title} ／ APL {VERSION} ／ ")
        _set_run_font(run, size=8, color=GRAY)
        fld1 = OxmlElement("w:fldChar")
        fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        r2 = p.add_run()
        r2._r.append(fld1)
        r2._r.append(instr)
        r2._r.append(fld2)
        _set_run_font(r2, size=8, color=GRAY)


def cover(doc: Document, title: str, subtitle: str, purpose: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run("CSV Tool ／ データ集約")
    _set_run_font(r, size=14, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(title)
    _set_run_font(r, size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    _set_run_font(r, size=12, color=TEAL)

    meta = [
        ("対象版", f"APL {VERSION}"),
        ("作成日", DOC_DATE),
        ("文書種別", "操作説明書（利用者向け）"),
        ("画像", "本文の【画像挿入】枠に、実画面のスクリーンショットを貼り付ける。"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, (k, v) in enumerate(meta):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = v
        _shade_cell(t.cell(i, 0), "E8EEF4")
        for j in range(2):
            _set_cell_border(t.cell(i, j))
            for para in t.cell(i, j).paragraphs:
                para.paragraph_format.space_after = Pt(2)
                if para.runs:
                    _set_run_font(para.runs[0], size=10, bold=(j == 0))

    h(doc, "本書の目的", 1)
    body(doc, purpose)
    note(
        doc,
        "画像について",
        "画面キャプチャは未掲載である。【画像挿入】枠の内側に、枠の指示どおりの画面を貼り付ける。"
        "枠は残し、下の図番号も残す。",
    )
    doc.add_page_break()


def h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        _set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 11.5, bold=True, color=NAVY)


def body(doc: Document, text: str, *, bold=False, size=10.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=bold, color=GRAY)


def bullets(doc: Document, items: list[str], *, numbered=False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for it in items:
        p = doc.add_paragraph(it, style=style)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            _set_run_font(r, size=10.5, color=GRAY)


def note(doc: Document, title: str, text: str, *, kind="注意") -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    _shade_cell(cell, HINT_BG)
    _set_cell_border(cell)
    p0 = cell.paragraphs[0]
    r = p0.add_run(f"【{kind}】{title}")
    _set_run_font(r, size=10, bold=True, color=RGBColor(0x7A, 0x4F, 0x01))
    p1 = cell.add_paragraph()
    r2 = p1.add_run(text)
    _set_run_font(r2, size=10, color=GRAY)
    doc.add_paragraph()


def img_slot(doc: Document, fig_id: str, caption: str, hint: str) -> None:
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0 = t.cell(0, 0)
    _shade_cell(c0, "FFF8E7")
    _set_cell_border(c0)
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"【画像挿入】{fig_id}")
    _set_run_font(r, size=11, bold=True, color=RGBColor(0x8A, 0x5A, 0x00))
    p2 = c0.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(hint)
    _set_run_font(r2, size=9.5, color=GRAY)
    p3 = c0.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("（この枠の内側にスクリーンショットを貼り付け）")
    _set_run_font(r3, size=9, color=RGBColor(0x99, 0x77, 0x22))
    c1 = t.cell(1, 0)
    _set_cell_border(c1)
    p4 = c1.paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"{fig_id}　{caption}")
    _set_run_font(r4, size=9.5, bold=True, color=NAVY)
    doc.add_paragraph()


def table(doc: Document, headers: list[str], rows: list[list[str]], *, col_cm: list[float] | None = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for j, htxt in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(htxt)
        _set_run_font(r, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(cell, HEAD_BG)
        _set_cell_border(cell)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            _set_run_font(r, size=9, color=GRAY)
            if i % 2 == 1:
                _shade_cell(cell, ALT_BG)
            _set_cell_border(cell)
    if col_cm:
        for row in t.rows:
            for j, w in enumerate(col_cm):
                row.cells[j].width = Cm(w)
    doc.add_paragraph()


def code_block(doc: Document, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    _shade_cell(cell, CODE_BG)
    _set_cell_border(cell)
    first = True
    for line in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(line if line else " ")
        _set_run_font(r, name="Consolas", size=9, color=RGBColor(0x22, 0x22, 0x22))
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "游ゴシック")
    doc.add_paragraph()


# ===========================================================================
# 1. シナリオ編集画面説明書
# ===========================================================================

def build_ui_manual() -> Document:
    doc = new_doc()
    add_footer(doc, "シナリオ編集画面説明書")
    cover(
        doc,
        "データ集約　シナリオ編集画面説明書",
        "画面の使い方と設定例　Ver1.1.10.6",
        "データ集約の「シナリオ編集」画面の使い方をまとめたものである。"
        "画面で何を選ぶと、結果がどうなるかを中心に説明する。",
    )

    h(doc, "1. この画面でできること", 1)
    body(doc, "マスタの各項目（列）について、「どのファイルの、どの場所から値を取るか」を決める画面である。")
    bullets(
        doc,
        [
            "開き方：メイン画面の項目行にある「編集」を押す。",
            "確定：「登録」を押したときだけ設定が残る。行を切り替えただけでは名前は確定しない。",
            "閉じる：「閉じる」は、登録していない変更を捨てて戻る。",
            "Undo：直前の削除または登録を、1回だけ元に戻せる。",
        ],
    )
    img_slot(doc, "図1-1", "メイン画面の「編集」ボタン", "項目一覧と「編集」ボタンが見える状態。")

    h(doc, "2. 画面の見方", 1)
    table(
        doc,
        ["場所", "見方"],
        [
            ["左上", "この項目の取得ルール一覧。上から順に処理する。"],
            ["左下", "選んだルールの内容（全文）。"],
            ["左下のボタン", "▲▼（順番）／追加／複製／削除／Undo／デバッグ"],
            ["右上", "項目名、種別、シナリオ名"],
            ["右の本文", "種別ごとの設定。整形（DSL）の右にある灰色の小ボタンからテスト画面を開く。"],
            ["右下", "登録／閉じる"],
        ],
    )
    img_slot(doc, "図2-1", "シナリオ編集画面の全体", "左に一覧、右に設定が見える全体。")

    h(doc, "2.1 一覧のボタン", 2)
    table(
        doc,
        ["ボタン", "結果"],
        [
            ["追加", "新しい取得ルールが末尾に増える。"],
            ["複製", "選んでいるルールのコピーが増える。少しだけ変えて使うときに便利である。"],
            ["削除", "選んでいるルールが消える。メイン画面の要約もすぐ変わる。"],
            ["▲／▼", "処理する順番が上下に入れ替わる。"],
            ["デバッグ", "実際に取る値を試し見できる。マスタには書き込まない。"],
        ],
    )
    note(
        doc,
        "種別の混ぜ合わせ",
        "同じ項目の中で、「セル座標から取得」と「名前から取得」は同時に使えない。別の項目に分ける。",
    )

    h(doc, "2.2 入力のコツ（V1.1.10.6）", 2)
    bullets(
        doc,
        [
            "セル座標は、小文字で打ってもすぐ大文字になる（例：b5 → B5）。固定値の文字はそのままである。",
            "数値やリストは、選んでいるときだけマウスホイールで変わる。誤って回しても変わりにくい。",
        ],
    )

    h(doc, "3. セル座標から取得", 1)
    body(doc, "Excel のセルから値を取るときの設定である。上から、ファイル → シート → 主キー → 連携キー → 結合キー、の順に決める。")
    img_slot(doc, "図3-1", "セル座標から取得の設定全体", "右ペインで 1〜5 の見出しが見える状態。")

    h(doc, "3.1 ファイルの絞り込み", 2)
    table(
        doc,
        ["設定", "結果"],
        [
            ["検索条件＝含む、ファイル名＝空", "候補のすべてのファイルが対象になる。"],
            ["含む、ファイル名＝光特性,紐づけ", "名前に「光特性」または「紐づけ」が付くファイルだけが対象になる。"],
            ["含まない、ファイル名＝一時", "「一時」が付くファイルは対象外になる。"],
            ["完全一致、ファイル名＝一覧表", "名前が「一覧表」のファイルだけが対象になる（拡張子は見ない）。"],
            ["ファイル種別", "チェックした種類だけが対象になる。CSV は最初はオフである。"],
        ],
    )
    img_slot(doc, "図3-2", "ファイル名の複数指定", "検索条件「含む」、ファイル名「光特性,紐づけ」。")

    h(doc, "3.2 シートの絞り込み", 2)
    table(
        doc,
        ["設定", "結果"],
        [
            ["左端シート", "ブックの一番左のシートだけを見る。シート名の入力は使わない。"],
            ["含む、シート名＝R_,実装", "名前に「R_」または「実装」が付くシートを見る。"],
            ["含まない", "書いた文字が付くシートは見ない。"],
            ["空欄（左端以外）", "該当するシートがない。"],
        ],
    )
    body(doc, "CSV にはシートがないため、この設定は効かない。シート名にカンマが含まれると、うまく分かれない。")
    img_slot(doc, "図3-3", "シート名の複数指定", "条件「含む」、シート名「R_,実装」。")

    h(doc, "3.3 主キー（どこから、何件取るか）", 2)
    body(doc, "主キーは、「何件の行を作るか」と「その項目の値」を決める。")
    table(
        doc,
        ["設定", "結果"],
        [
            ["セル座標＝F4、行オフセット＝1、空白まで", "F4、F5、F6…と下へ進み、空欄で止まる。"],
            ["行オフセット＝0、列オフセット＝0、N件＝1", "そのセルを 1 個だけ取る（1ファイルで1件）。"],
            ["終端", "シートの終わりまで取る。"],
            ["N件＝10", "最大 10 件で止める。"],
        ],
    )
    note(
        doc,
        "取得件数の上限",
        "「N件」に達すると、その先にデータがあってもそこで止める。"
        "上限の見直しはデバッグで確認する（警告が出る）。本番の一括実行では警告画面は出さず、上限件数のまま完了する。",
    )
    note(
        doc,
        "0/0 のとき",
        "行も列も動かさないときは、「空白まで」「終端」は使えない。件数（N件）を指定する。",
    )

    h(doc, "3.3.1 主キーのオプション", 3)
    table(
        doc,
        ["設定", "結果"],
        [
            ["非表示・フィルタ行を除く＝ON", "Excel で隠している行（保存済み）は取らない。見える行だけが結果になる。CSV では効かない。"],
            ["主キーをスキップ＝ON、照合が空", "主キーが空の行は結果に残さない。"],
            ["照合＝,-", "空欄と「-」の行は結果に残さない。"],
            ["スキップ行を前置に使う＝ON", "落とした行の連携値（前回値を保持している項目）を、次の行以降の穴埋めに使う。"],
        ],
    )
    img_slot(doc, "図3-4", "主キーの設定", "座標・オフセット・終わり方・スキップ・加工・整形が見える状態。")

    h(doc, "3.3.2 加工と整形", 3)
    body(doc, "取った文字の直し方は、次の順である。")
    bullets(
        doc,
        [
            "セル内の改行は先に消える（「電」と「源」が分かれていても「電源」になる）。",
            "加工（チェック）：トリム、全角→半角、年月日変換。付けたものだけ効く。",
            "整形（DSL）：左から順にコマンドを当てる。書き方は別紙の整形 DSL 説明を見る。入力欄の右の灰色ボタンで、当てた結果をその場で確認できる（5章）。",
        ],
        numbered=True,
    )

    h(doc, "3.3.3 書込みモード", 3)
    table(
        doc,
        ["選び方", "結果"],
        [
            ["空き上書き", "行き先が空のときだけ書く。"],
            ["強制上書き", "行き先に値があっても上書きする。"],
            ["行追加", "新しい行として足す。"],
            ["複写追加", "同じキーの行を増やして書く。"],
        ],
    )

    h(doc, "3.4 連携キー（同じ行の別の値）", 2)
    body(doc, "主キーと同じ回数・同じ段で、別のセル（または決まった文字）を取り、別の項目へ入れる。終わり方と書込み方は主キーに従う。")
    table(
        doc,
        ["設定", "結果"],
        [
            ["セル座標＝H4、行オフセット＝1", "主キーが F4,F5,F6 なら、品名は H4,H5,H6 になる。"],
            ["セル座標＝H4+I4", "H と I の文字を、間に何も入れずにつなげる（空のセルは空文字）。"],
            ["値種別＝固定値、文字＝PAST", "毎回「PAST」が入る。"],
            ["空欄は前回値を保持＝ON", "同じシートの中で、空欄には直前の値を入れる。ファイルやシートが変わるとリセットする。"],
        ],
    )
    body(doc, "連携項目に選んだ項目は、自分ではシナリオを持てない（一覧が灰色になる）。")
    bullets(
        doc,
        [
            "まだ 0 件のときは「＋ 連携キー追加」で 1 件目を作る。",
            "1 件以上あるときは、各ブロックの「下に挿入」で、その直後に足す。",
            "「削除」でその定義を消す。上限は 50 件である。",
        ],
    )
    img_slot(doc, "図3-5", "連携キーの設定", "座標、連携項目、「空欄は前回値を保持」、「下に挿入」が見える状態。")

    h(doc, "3.5 結合キー（別の項目と行をそろえる）", 2)
    body(doc, "別ファイルや別項目の行を、同じ値で 1 行にそろえるための目印である。複数あるときは、すべて一致した行だけがそろう。")
    table(
        doc,
        ["設定", "結果"],
        [
            ["セル座標＝A5、行オフセット＝1、結合項目＝機器番号", "主キーと同じ段の A 列を取り、機器番号が同じ行とつなぐ。"],
            ["定義を 2 つ（機器番号と日付）", "両方同じ行だけがそろう。"],
        ],
    )
    body(doc, "結合キーの座標は 1 セルだけである。「H4+I4」のようなつなぎ方は使えない。")
    img_slot(doc, "図3-6", "結合キーの設定", "セル座標と結合項目が見える状態。")

    h(doc, "4. 名前から取得", 1)
    body(doc, "フォルダ名やファイル名から文字を取り、すでにできた行へ足す。新しい行は増やさない。")
    table(
        doc,
        ["設定", "結果"],
        [
            ["検索対象＝ファイル名、区切文字＝_、取得ブロック＝2", "「2024_LOT12_xxx.xlsx」なら「LOT12」が入る。"],
            ["関連付け項目＝先頭の項目", "同じファイル（同じパス）の行へ入る。"],
            ["空き上書き", "そのマスが空のときだけ入る。"],
        ],
    )
    img_slot(doc, "図4-1", "名前から取得の設定", "検索・抜取り・関連付けが見える状態。")

    h(doc, "5. 整形 DSL のテスト", 1)
    body(
        doc,
        "整形（DSL）の入力欄の右にある灰色の小さなボタンを押すと、「DSLテスト」画面が開く。"
        "マスタやシナリオにはすぐには書き込まない。結果をシナリオへ反映するときは「ペースト」を使う。",
    )
    bullets(
        doc,
        [
            "主キー・連携キー・結合キー・名前から取得の、それぞれの整形欄から開ける。",
            "同時に開くのは 1 つまで。別の整形欄のボタンを押すと、前のテスト画面は閉じる。",
            "テスト画面は、灰色ボタンの下に出す。画面下にはみ出すときは、ボタンの上側に出す。",
            "右ペインの別の入力へ移る、シナリオ行を切り替える、シナリオ編集を閉じる、とテスト画面も閉じる。",
        ],
    )
    img_slot(
        doc,
        "図5-1",
        "整形（DSL）入力と灰色のテストボタン",
        "シナリオ編集の右ペイン。整形（DSL）入力欄の右端に、灰色の小さな四角ボタンが見える状態。ボタンにマウスを乗せたツールチップが出ていればなお良い。",
    )

    h(doc, "5.1 画面の見方", 2)
    table(
        doc,
        ["場所", "見方"],
        [
            ["冒頭の説明", "テスト結果は参考であること、シナリオへ入れるときはペーストすること。"],
            ["DSLテスト用文字列", "整形の対象にする試し文字。右の灰色ボタンで、あらかじめ決めた規定値に戻せる。"],
            ["DSLコマンド入力", "試すコマンド列。開いたときの整形欄の内容がコピーされる。1 行で入力する。"],
            ["DSLコマンド実行表示", "ステップまたは一括で、いま当てたコマンド（書いたとおりの形）を表示する。書き換えられない。"],
            ["DSL結果", "当てたあとの文字。書き換えられない。右の灰色ボタンで、実行表示と結果を空にする。"],
            ["ステップ／一括実行／ペースト", "試し実行と、シナリオ側への反映。"],
            ["閉じる", "テスト画面を閉じる。試し文字は、Excel を終了するまで次のテストでも使う。"],
            ["右側", "コマンドの説明表（整形欄のツールチップと同じ内容）。"],
        ],
    )
    img_slot(
        doc,
        "図5-2",
        "DSLテスト画面の全体",
        "左右に分かれた DSLテスト画面の全体。左に入力・実行・結果とボタン、右にコマンド説明が見える状態。冒頭の説明文が欠けていないこと。",
    )

    h(doc, "5.2 試し文字の覚え方", 2)
    bullets(
        doc,
        [
            "Excel を起動して最初に開いたときは、あらかじめ決めた規定値が入る。",
            "閉じたあとも、空欄を含めて同じ試し文字を使う（Excel を終了するまで）。",
            "規定値に戻したいときは、試し文字の右の灰色ボタンを押す。",
        ],
    )

    h(doc, "5.3 ステップと一括実行", 2)
    table(
        doc,
        ["操作", "結果"],
        [
            ["ステップ", "コマンドを左から 1 つずつ当てる。実行表示には、いままで当てた分が出る。"],
            ["最終コマンドの次のステップ", "初期状態になる（実行表示と結果が空。クリアボタンと同じ）。"],
            ["その次のステップ", "再び 1 コマンド目から当てる。"],
            ["一括実行", "すべてのコマンドを一度に当てる。実行表示には全コマンドが出る。"],
            ["コマンドが空", "実行表示に「(コマンドなし)」、結果は試し文字のまま。"],
        ],
    )
    img_slot(
        doc,
        "図5-3",
        "ステップ実行の途中",
        "DSLコマンド入力に複数コマンドがあり、実行表示には途中までのコマンド、結果には変換後の文字が見える状態。",
    )

    h(doc, "5.4 ペーストと構文エラー", 2)
    bullets(
        doc,
        [
            "ペースト：構文が正しければ、開いた元の整形欄へコマンドを書き込む。テスト画面は閉じない。",
            "構文が正しくないとき：該当するコマンドが赤の太字になり、警告の画面が出る。シナリオ側へは書き込まない。",
            "ステップ／一括実行でも、構文が正しくないときは同じ警告になる。",
        ],
    )
    img_slot(
        doc,
        "図5-4",
        "構文エラー時の表示",
        "DSLコマンド入力の誤り箇所が朱書き太字で、警告の小さな画面（「構文エラー」）が重なっている状態。入力欄の背景は赤くしない。",
    )
    note(
        doc,
        "テスト結果は参考",
        "テスト画面の結果は、その場の確認用である。シナリオに残すには「ペースト」のあと、シナリオ編集の「登録」が必要である。",
    )

    h(doc, "6. デバッグ", 1)
    body(doc, "「デバッグ」を押すと、今の設定で実際に取る値を見られる。マスタには書かない。")
    bullets(
        doc,
        [
            "前回値を保持している項目は、結果の名前の先頭に「・」が付く。",
            "おかしいときは、ファイル・シート・座標・オフセットの順に見直す。",
            "主キーの取得件数が「N件」の上限に達したときは、分かりやすい警告を出す。"
            "シナリオの取得件数を見直し、必要なら直してから本番の一括へ進む。",
        ],
    )
    img_slot(doc, "図6-1", "デバッグの結果", "結果一覧。可能なら先頭に「・」が付いた列が見える状態。")

    h(doc, "7. よく使う設定例", 1)
    h(doc, "7.1 一覧を下へ取る", 2)
    table(
        doc,
        ["項目", "値"],
        [
            ["ファイル名", "空、または「履歴」を含む"],
            ["シート", "含む／R_"],
            ["主キー", "F4、行オフセット 1、空白まで、行追加"],
        ],
    )
    body(doc, "結果：F4 から下の値が、空欄の手前まで行になる。")

    h(doc, "7.2 隣のセルも一緒に取る", 2)
    table(
        doc,
        ["項目", "値"],
        [
            ["主キー", "F4、行オフセット 1"],
            ["連携キー", "H4（または H4+I4）、行オフセット 1、連携項目＝品名"],
        ],
    )
    body(doc, "結果：機器番号と品名が、同じ行のセットで入る。")

    h(doc, "7.3 見出し行を落とし、見出しの品名を明細へ回す", 2)
    table(
        doc,
        ["項目", "値"],
        [
            ["主キーをスキップ", "ON（照合は空、または ,- ）"],
            ["スキップ行を前置に使う", "ON"],
            ["連携の空欄は前回値を保持", "ON"],
        ],
    )
    body(doc, "結果：見出し行は残らない。明細の空の品名には、見出しの品名が入る。")

    h(doc, "7.4 2つのファイルを機器番号でつなぐ", 2)
    table(
        doc,
        ["シナリオ", "ファイル名", "主キー", "結合キー"],
        [
            ["光特性", "光特性を含む", "測定値のセル", "機器番号のセル → 項目「機器番号」"],
            ["紐づけ", "紐づけを含む", "製番のセル", "機器番号のセル → 項目「機器番号」"],
        ],
    )
    body(doc, "結果：機器番号が同じ行に、測定値と製番が横に並ぶ。")

    h(doc, "8. メイン画面のスキャンと一括（概要）", 1)
    body(
        doc,
        "シナリオ編集の前後で使う、メイン画面の動きである。"
        "基準フォルダのスキャンと、一括実行時の見え方をまとめる。",
    )

    h(doc, "8.1 基準フォルダのスキャン", 2)
    table(
        doc,
        ["状態", "ボタンの動き"],
        [
            ["スキャン中", "一括実行・読込・保存・デバッグ・出力などを押せない。"],
            ["スキャン成功", "通常どおり使える。一括では、この結果を引き継ぎ、同じ条件の再スキャンを省略する。"],
            ["スキャン失敗（自動リトライ後）", "一括実行とデバッグだけ押せない。読込・保存・出力・再検索は使える。"],
        ],
    )
    bullets(
        doc,
        [
            "失敗時は自動でやり直す（既定 3 回・間隔 約 1 秒）。それでもだめなときは、条件を直して「検索実行」する。",
            "再検索が成功すると、一括実行とデバッグも再び使える。",
        ],
    )

    h(doc, "8.2 一括実行の進捗", 2)
    bullets(
        doc,
        [
            "ネットワーク上のフォルダ（UNC など）は、いったん TEMP へマウントしてから読み取る（並列コピー）。進捗には「マウント」と出る。",
            "進捗の件数（N/M）は、並列でも完了した分だけが増える（途中で減らない）。",
            "本番の一括では、主キー取得件数の上限超過でも警告画面は出さない。上限の確認はデバッグで行う。",
        ],
    )

    h(doc, "9. ほかの説明書", 1)
    table(
        doc,
        ["文書", "内容"],
        [
            ["主キー・連携キー・結合キーの動作概念_V11106", "3つのキーの違いと例"],
            ["シナリオファイルJson構成と概要_V11106", "保存ファイルの見方"],
            ["整形DSL_コマンドリファレンス_V11106", "文字の直し方と DSLテスト画面"],
        ],
    )
    return doc


# ===========================================================================
# 2. キー動作概念
# ===========================================================================

def build_keys_concept() -> Document:
    doc = new_doc()
    add_footer(doc, "主キー・連携キー・結合キーの動作概念")
    cover(
        doc,
        "データ集約　主キー・連携キー・結合キーの動作概念",
        "3つのキーの違いと設定例　Ver1.1.10.6",
        "主キー・連携キー・結合キーの役割と、こう設定するとこう並ぶ、という例をまとめたものである。",
    )

    h(doc, "1. 3つのキーの違い", 1)
    table(
        doc,
        ["キー", "役割", "こう設定すると", "こうなる"],
        [
            ["主キー", "何件取るか、その項目の値", "F4 から下へ、空白まで", "空欄の手前まで行ができる"],
            ["連携キー", "同じ行の別の値", "H4、行オフセットは主キーと同じ", "各行に品名などが付く"],
            ["結合キー", "別の表と行をそろえる目印", "両方のシナリオで結合項目＝機器番号", "機器番号が同じ行が 1 行になる"],
        ],
    )
    body(doc, "連携キーは「同じファイルの隣のセル」である。結合キーは「別のファイルや別の項目とつなぐ目印」である。混ぜて考えない。")
    img_slot(
        doc,
        "図1-1",
        "3つのキーのイメージ",
        "Excel の列と、結果表の列を矢印で結んだ図。主キー＝件数、連携＝同じ行、結合＝つなぎ、と書き分ける。",
    )

    h(doc, "2. 基本のルール", 1)
    bullets(
        doc,
        [
            "主キーで 3 件取ったら、連携も結合も 3 回見る。足りない分は空欄になる（前の値の使い回しはしない）。",
            "進み方（オフセット）は、キーごとに別々に決められる。",
            "同じ行の隣を取りたいときは、行の進みを同じ数にする（多くは 1）。",
            "毎回同じセルを使いたいときは、そのキーの進みを 0 にする。",
        ],
    )

    h(doc, "3. 進み方（オフセット）", 1)
    body(doc, "基準のセルから、「何行・何列ずつ進むか」である。")
    table(
        doc,
        ["回", "主キー F4・行1", "連携 H4・行1", "結合 A4・行1"],
        [
            ["1回目", "F4", "H4", "A4"],
            ["2回目", "F5", "H5", "A5"],
            ["3回目", "F6", "H6", "A6"],
        ],
    )
    body(doc, "連携だけ進みを 0 にすると、毎回 H4 になる。")
    body(doc, "連携を「D10+D11」、行の進み 1 にすると、1回目は D10 と D11 をつなぎ、2回目は D11 と D12 をつなぐ。")
    img_slot(doc, "図3-1", "セルの進み方", "F4／H4／A4 を色分けし、下へ矢印を付けた図。")

    h(doc, "4. 行を飛ばしたときの動き", 1)
    body(doc, "主キーで行を落とすと、結果の行数は減る。ただし連携と結合は、落としたあとも「元の表の同じ段」を見る。ずれない。")
    table(
        doc,
        ["元の表", "設定", "結果"],
        [
            ["4行目が見える、5行目が隠れている、6行目が見える", "非表示行を除く＝ON", "4行目と6行目の 2 件。連携も 4 行目と 6 行目を見る"],
            ["4行目の主キーが空、5・6行目に番号", "主キーをスキップ＝ON", "5・6行目の 2 件"],
        ],
    )
    note(
        doc,
        "非表示の見方",
        "ファイルに保存されている「隠し」だけが対象である。画面でフィルタを変えても、保存し直すまで結果は変わらない。",
    )

    h(doc, "5. 空欄を前の値で埋める", 1)
    body(doc, "連携キーの「空欄は前回値を保持」を ON にすると、同じシートの中で、空欄に直前の値が入る。先頭が空なら空のままである。シートやファイルが変わると消える。")
    body(doc, "見出し行を主キーから落とすときは、次の 2 つを両方 ON にする。")
    bullets(
        doc,
        [
            "主キーをスキップ",
            "スキップ行を前置に使う（落とした行の品名などを、次の行の穴埋めに使う）",
        ],
    )
    img_slot(
        doc,
        "図5-1",
        "見出しを落として品名を回すイメージ",
        "見出し行（主キー空・品名あり）と明細（主キーあり・品名空）。結果の品名が見出しと同じになる図。",
    )

    h(doc, "6. 設定例", 1)
    h(doc, "6.1 同じシートで 3 つそろえる", 2)
    table(
        doc,
        ["Excel", "A ロット", "F 機器番号", "H 品名"],
        [
            ["4行目", "L-01", "SN-100", "電源A"],
            ["5行目", "L-01", "SN-101", "電源A"],
            ["6行目", "L-02", "SN-200", "ファン"],
        ],
    )
    table(
        doc,
        ["キー", "座標", "行の進み", "行き先"],
        [
            ["主キー", "F4", "1", "機器番号"],
            ["連携", "H4", "1", "品名"],
            ["結合", "A4", "1", "ロットで他とつなぐとき"],
        ],
    )
    body(doc, "結果：3 行でき、各行に機器番号・品名・ロットがそろう。")

    h(doc, "6.2 2つのファイルを機器番号でつなぐ", 2)
    table(
        doc,
        ["ファイル", "取りたい値", "結合の目印"],
        [
            ["光特性履歴", "測定値", "同じ行の機器番号"],
            ["紐づけ履歴", "製番", "同じ行の機器番号"],
        ],
    )
    body(doc, "結果：機器番号が同じ行に、測定値と製番が横に並ぶ。")
    img_slot(doc, "図6-1", "2ファイルをつなぐイメージ", "左：光特性。右：紐づけ。中央：機器番号で 1 行になる図。")

    h(doc, "6.3 見出しスキップと穴埋め", 2)
    table(
        doc,
        ["元の行", "主キー", "品名", "結果"],
        [
            ["見出し", "空", "電源ユニット", "行は残らない。品名だけ次へ回る"],
            ["明細", "SN-100", "空", "1行残る。品名＝電源ユニット"],
            ["明細", "SN-101", "空", "1行残る。品名＝電源ユニット"],
            ["見出し", "-", "ファン", "行は残らない（照合に「-」があるとき）"],
            ["明細", "SN-200", "空", "1行残る。品名＝ファン"],
        ],
    )
    body(doc, "必要なスイッチ：主キーをスキップ、照合「,-」、スキップ行を前置に使う、連携の前回値保持。")

    h(doc, "6.4 ファイルの区別を固定文字で付ける", 2)
    body(doc, "同じ項目にルールを 2 本置き、過去ファイル側だけ連携を固定値「PAST」にする。")
    body(doc, "結果：過去から来た行に「PAST」が付く。")

    h(doc, "7. うまくいかないとき", 1)
    table(
        doc,
        ["見た目", "よくある原因", "直し方"],
        [
            ["連携が 1 行ずれている", "行の進みが主キーと違う", "両方とも同じ数（多くは 1）にする"],
            ["見出しの品名が入らない", "スキップ行を前置に使う、がオフ", "スキップと前回値保持を両方 ON にする"],
            ["ファイルがつながらない", "結合の値が左右で違う、または見るセルが違う", "デバッグで結合の値を見比べる"],
            ["項目が灰色で触れない", "ほかのルールの連携先になっている", "先にそちらの連携を外す"],
            ["整形の結果がおかしい", "コマンドの順や引数", "整形欄の灰色ボタンから DSLテストで 1 つずつ確認する"],
            ["件数が想定より少ない", "主キーの「N件」上限に達している", "デバッグで警告を確認し、取得件数を見直す"],
            ["デバッグでは警告、一括では出ない", "本番一括は上限超過でも警告画面を出さない仕様", "上限の確認はデバッグ側で行う"],
        ],
    )
    return doc


# ===========================================================================
# 3. JSON 構成（操作者向け：保存ファイルの見方）
# ===========================================================================

def build_json_doc() -> Document:
    doc = new_doc()
    add_footer(doc, "シナリオファイルJson構成と概要")
    cover(
        doc,
        "データ集約　シナリオファイル Json 構成と概要",
        "保存ファイルの見方　Ver1.1.10.6",
        "「シナリオ保存」でできるファイルの見方である。"
        "画面のどこが、ファイルのどこに残るかを中心に説明する。普段は画面から編集すれば足りる。",
    )

    h(doc, "1. このファイルは何か", 1)
    body(doc, "項目の並び、各項目の取り方、探すフォルダ、Excel への出し方を、ひとまとめにした設定ファイルである。")
    bullets(
        doc,
        [
            "メインの「シナリオ保存」で書き出す。「シナリオ読込」で戻す。",
            "文字コードは UTF-8 である。",
            "結果にパスやファイル名の列を付ける設定は、このファイルには残らない（実行のたびに画面で選ぶ）。",
            "一括へ渡すスキャン結果（見つかったファイル一覧）も、このファイルには残らない。"
            "画面で検索した直後の一括では再スキャンを省略するが、保存・読込の対象外である。",
        ],
    )
    img_slot(doc, "図1-1", "シナリオ保存と、保存したファイルの先頭", "左：保存ボタン。右：ファイルの先頭が見える状態。")

    h(doc, "2. 大きな枠", 1)
    table(
        doc,
        ["名前", "画面での意味"],
        [
            ["items", "項目の一覧。並びが処理の順である。"],
            ["scan", "基準フォルダ、サブフォルダ、拡張子、キーワード。"],
            ["excel_options", "Excel タブ（新規シート、並べ替えなど）。"],
            ["match_keys", "行をそろえる項目の指定（使う場合）。"],
        ],
    )

    h(doc, "3. 項目（items）", 1)
    table(
        doc,
        ["名前", "画面での意味"],
        [
            ["name", "項目名（列の見出し）。"],
            ["id", "内部の番号。通常は触らない。"],
            ["sources", "その項目の取得ルール（1 項目に複数可）。"],
            ["write_mode", "項目全体の書込み。実際の正は、各ルール側の書込みである。"],
        ],
    )
    note(
        doc,
        "種別の混ぜ合わせ",
        "1つの項目の中で、セルから取るルールと、名前から取るルールは同時に置けない。",
    )

    h(doc, "4. セルから取るルール", 1)
    body(doc, "画面の「セル座標から取得」に対応する。")
    table(
        doc,
        ["画面", "ファイル上の名前", "例"],
        [
            ["セル座標", "cell_ref", "F4"],
            ["行／列の進み", "row_offset / col_offset", "1 / 0"],
            ["空白まで", "repeat_until_empty が true", "下へ空欄まで"],
            ["終端", "repeat_until_last が true", "終わりまで"],
            ["N件", "repeat_max", "10"],
            ["主キーをスキップ", "skip_empty_primary", "true"],
            ["スキップの照合", "skip_primary_match", ",-"],
            ["スキップ行を前置に使う", "skip_carry_seed", "true"],
            ["非表示行を除く", "skip_hidden_rows", "true"],
            ["シート名", "sheet_name", "R_"],
        ],
    )

    h(doc, "4.1 画面の細かい設定（ui_scenario_source_v1）", 2)
    table(
        doc,
        ["画面", "ファイル上の名前", "例"],
        [
            ["ファイル名", "file_pattern", "光特性,紐づけ"],
            ["検索条件", "file_name_rule", "含む"],
            ["ファイル種別", "ext_checked", ".xlsx など"],
            ["シート条件", "sheet_rule", "含む"],
            ["加工", "cell_checks", "トリム"],
            ["整形（DSL）", "value_shape_script", "trim"],
            ["書込み", "write_mode_cell_key", "append（行追加）"],
            ["連携キー", "link_defs", "配列"],
            ["結合キー", "join_defs", "配列"],
        ],
    )

    h(doc, "4.2 連携キー（link_defs）", 2)
    table(
        doc,
        ["画面", "ファイル上の名前", "例"],
        [
            ["連携項目", "item", "品名"],
            ["値種別", "mode", "セル座標 または 固定値"],
            ["座標／固定値", "cell", "H4 または H4+I4 または PAST"],
            ["行／列の進み", "row / col", "1 / 0"],
            ["加工", "checks", "トリム"],
            ["整形", "value_shape_script", "trim"],
            ["空欄は前回値を保持", "carry_empty", "true"],
        ],
    )

    h(doc, "4.3 結合キー（join_defs）", 2)
    table(
        doc,
        ["画面", "ファイル上の名前", "例"],
        [
            ["結合項目", "item", "機器番号"],
            ["セル座標", "cell", "A5（1セルだけ）"],
            ["行／列の進み", "row / col", "1 / 0"],
        ],
    )

    h(doc, "5. 名前から取るルール", 1)
    table(
        doc,
        ["画面", "ファイル上の名前", "例"],
        [
            ["フォルダ名／ファイル名", "source_type", "dir_name / file_name"],
            ["検索条件", "search_condition", "include（含む）"],
            ["検索文字", "search_text", "2024"],
            ["区切文字", "delimiter", "_"],
            ["取得ブロック", "part_index", "2"],
            ["関連付け項目", "path_item", "機器番号"],
            ["書込み", "write_mode_name_key", "fill_in（空き上書き）"],
            ["整形", "value_shape_script", "trim"],
        ],
    )

    h(doc, "6. 短い例", 1)
    body(doc, "機器番号を F4 から下へ取る、いちばん簡単な形である。")
    code_block(
        doc,
        """{
  "version": 1,
  "items": [
    {
      "id": "item_0",
      "name": "機器番号",
      "sources": [
        {
          "type": "cell",
          "sheet_name": "R_",
          "cell_ref": "F4",
          "row_offset": 1,
          "col_offset": 0,
          "repeat_until_empty": true,
          "skip_empty_primary": false,
          "skip_hidden_rows": false,
          "ui_scenario_source_v1": {
            "file_pattern": "光特性",
            "file_name_rule": "含む",
            "sheet_rule": "含む",
            "value_shape_script": "trim",
            "write_mode_cell_key": "append",
            "link_defs": [],
            "join_defs": []
          }
        }
      ]
    }
  ],
  "scan": {
    "start_path": "D:\\\\data",
    "recursive": true,
    "extensions": [".xlsx", ".xls"],
    "keyword": ""
  }
}""",
    )
    img_slot(doc, "図6-1", "実ファイルの連携・結合のあたり", "保存した JSON で link_defs と join_defs が見える範囲。")

    h(doc, "7. 直すときの注意", 1)
    bullets(
        doc,
        [
            "連携項目・結合項目が空の定義は、画面から登録しても残らない。",
            "整形の書き方が間違っていると、登録できない。DSLテスト画面の試し文字は、このファイルには残らない。",
            "行も列も進みが 0 のとき、「空白まで」「終端」は使えない。",
            "普段は画面で直し、「シナリオ保存」する。手で直すときはバックアップを残す。",
        ],
    )
    return doc


# ===========================================================================
# 4. 整形 DSL
# ===========================================================================

def build_dsl_doc() -> Document:
    doc = new_doc()
    add_footer(doc, "整形DSLコマンドリファレンス")
    cover(
        doc,
        "データ集約　整形 DSL　コマンドリファレンス",
        "文字の直し方　Ver1.1.10.6",
        "取った文字を整えるときの書き方である。"
        "こう書くと、こう変わる、を一覧にする。",
    )

    h(doc, "1. いつ効くか", 1)
    body(doc, "加工（チェック）のあと、左から順に効く。主キー・連携キー・名前から取得で使える。")
    img_slot(doc, "図1-1", "整形（DSL）の入力欄", "主キーブロックの「整形（DSL）」と、右端の灰色テストボタンが見える状態。")

    h(doc, "2. 書き方", 1)
    bullets(
        doc,
        [
            "コマンドを左から順に当てる。",
            "区切りはカンマ（,）またはセミコロン（;）である。どちらでも同じである。",
            "文字の中にカンマがあるときは、\" \" で囲む。\" 自体は \"\" と書く。",
            "コマンド名の大文字・小文字はどちらでもよい。",
        ],
    )
    code_block(doc, 'trim,rep,"旧","新",case,upper\ntrim; left,3')

    h(doc, "3. コマンド一覧", 1)
    table(
        doc,
        ["コマンド", "何をするか", "書き方", "例と結果"],
        [
            ["trim", "前後の空白を消す", "trim", "「  AB  」→「AB」"],
            ["split", "指定した行だけ残す（1行目から数える）", "split,行番号", "「A（改行）B」に split,2 →「B」"],
            ["left", "左から n 文字", "left,数 または 式", "「ABCDEF」に left,3 →「ABC」"],
            ["right", "右から n 文字", "right,数 または 式", "「ABCDEF」に right,3 →「DEF」"],
            ["rep", "見つかった文字をすべて置き換える", "rep,元,先", "「aa-aa」に rep,\"aa\",\"X\" →「X-X」"],
            ["mid", "途中から指定の長さを切り出す", "mid,開始,長さ", "「ABCDEF」に mid,2,3 →「BCD」"],
            ["cut", "途中から指定の長さを消す", "cut,開始,長さ", "「ABCDEF」に cut,2,3 →「AEF」"],
            ["ins", "指定位置の直前に文字を入れる", "ins,位置,文字", "「ABCDEF」に ins,3,\"XY\" →「ABXYCDEF」"],
            ["padr", "右を指定文字で埋めて幅をそろえる", "padr,幅,文字", "「AB」に padr,5,\"0\" →「AB000」"],
            ["padl", "左を埋めて幅をそろえる", "padl,幅,文字", "「AB」に padl,5,\"0\" →「000AB」"],
            ["case", "大文字または小文字にする", "case,upper または case,lower", "「AbC」に case,upper →「ABC」"],
            ["wide", "全角を半角などにそろえる", "wide", "「ＡＢ１２」→「AB12」"],
            ["date", "日付を YYYY/MM/DD にする", "date", "「2026-09-01」→「2026/09/01」"],
        ],
    )
    body(doc, "位置は 1 から数える。left / right / mid / cut / ins は、数の代わりに次の式も使える。split と pad は数字だけである。")

    h(doc, "4. 位置の式（pos / len）", 1)
    table(
        doc,
        ["書き方", "意味"],
        [
            ["3", "3 文字、または 3 文字目"],
            ["pos(\"GH\")", "「GH」が始まる位置（1 から数える）。無いときは、そのコマンドは何もしない"],
            ["len()", "今の文字の長さ"],
            ["len(\"AB\")", "2（書いた文字の長さ）"],
            ["+ と - と ( )", "足す・引く・まとめる"],
        ],
    )
    body(doc, "元の文字を「ABCDEFGHIJK」（11文字）としたとき。")
    table(
        doc,
        ["書き方", "結果", "見方"],
        [
            ["left,pos(\"GH\")", "ABCDEFG", "G まで残す"],
            ["left,pos(\"GH\")-1", "ABCDEF", "G の手前まで残す"],
            ["ins,pos(\"G\"),\"123\"", "ABCDEF123GHIJK", "G の前に入れる"],
            ["ins,pos(\"G\")+1,\"123\"", "ABCDEFG123HIJK", "G の後ろに入れる"],
            ["left,pos(\"ZZ\")", "ABCDEFGHIJK", "見つからないので、そのまま"],
        ],
    )
    img_slot(
        doc,
        "図4-1",
        "位置の数え方",
        "ABCDEFGHIJK に 1〜11 の番号を付け、pos(\"GH\")＝7 と left,pos(\"GH\")-1 → ABCDEF を示した図。",
    )

    h(doc, "5. つなぎ書きの例", 1)
    table(
        doc,
        ["書き方", "元", "結果"],
        [
            ["trim; left,pos(\"-\")-1", "  AB-99  ", "AB"],
            ["rep,\"-\",\"/\"; left,pos(\"/\")-1", "2026-09-01", "2026"],
            ["wide; case,upper; left,3", "ａｂｃｄ", "ABC"],
            ["mid,pos(\"(\")+1,pos(\")\")-pos(\"(\")-1", "電源(12V)", "12V"],
        ],
    )

    h(doc, "6. よく使う型", 1)
    table(
        doc,
        ["目的", "書き方", "元 → 結果"],
        [
            ["ハイフンの前だけ", "left,pos(\"-\")-1", "ODN-164 → ODN"],
            ["括弧の中", "mid,pos(\"(\")+1,pos(\")\")-pos(\"(\")-1", "電源(12V) → 12V"],
            ["末尾 4 文字", "right,4", "SN00AB12 → AB12"],
            ["半角にしてゼロ埋め", "wide; padl,6,\"0\"", "１２３ → 000123"],
            ["日付をそろえる", "date", "20260901 → 2026/09/01"],
            ["ラベルを消す", "rep,\"機器:\",\"\"; trim", "機器: SN-1 → SN-1"],
            ["2行目だけ", "split,2; trim", "見出し（改行）値A → 値A"],
        ],
    )
    img_slot(doc, "図6-1", "整形の設定と結果", "左：整形の入力と灰色テストボタン。右：DSLテスト画面の結果。")

    h(doc, "7. 登録できないとき", 1)
    table(
        doc,
        ["メッセージの意味", "直し方"],
        [
            ["未知のコマンド", "一覧にある名前か確認する"],
            ["引数が不足", "必要な数だけ書く（例：rep は元と先の 2 つ）"],
            ["引数が不正", "式の括弧や \" が閉じているか見る"],
        ],
    )

    h(doc, "8. DSLテスト画面", 1)
    body(
        doc,
        "整形欄の右の灰色ボタンから開く確認用の画面である。"
        "コマンドを当てた結果を見てから、問題なければ「ペースト」で元の整形欄へ戻す。",
    )
    img_slot(
        doc,
        "図8-1",
        "DSLテスト画面の全体",
        "左：試し文字・コマンド・実行表示・結果とボタン。右：コマンド説明。冒頭説明が2行とも欠けていない状態。",
    )
    table(
        doc,
        ["欄／ボタン", "動き"],
        [
            ["DSLテスト用文字列", "当てる前の文字。Excel 起動後の初回は規定値。以降は空欄も含めて覚える。右の灰色ボタンで規定値に戻す。"],
            ["DSLコマンド入力", "試すコマンド。開いた整形欄の内容がコピーされる。1 行。"],
            ["DSLコマンド実行表示", "いま当てたコマンドの表示（書き換え不可）。背景は画面と同じ。"],
            ["DSL結果", "当てたあとの文字（書き換え不可）。背景は画面と同じ。右の灰色ボタンで実行表示と結果を空にする。"],
            ["ステップ", "左から 1 コマンドずつ当てる。最後の次は初期状態（両方空）、その次に 1 コマンド目へ戻る。"],
            ["一括実行", "全コマンドを一度に当てる。"],
            ["ペースト", "構文が正しければ元の整形欄へ書き込む。テスト画面は閉じない。"],
            ["閉じる", "画面を閉じる。試し文字は Excel 終了まで残る。"],
        ],
    )
    img_slot(
        doc,
        "図8-2",
        "ステップで途中まで当てたところ",
        "実行表示に trim,left,3 のように途中までのコマンド、結果に変換後の文字が見える状態。",
    )
    img_slot(
        doc,
        "図8-3",
        "構文エラー（朱書き太字）",
        "誤ったコマンド部分だけが赤の太字。警告画面「構文エラー」。入力欄全体の背景は赤くない。",
    )
    note(
        doc,
        "画面の出し方",
        "灰色ボタンのすぐ下に出す。画面の下端に収まらないときは、ボタンの上側に出す。"
        "詳細な画面の見方は「シナリオ編集画面説明書」の「5. 整形 DSL のテスト」も見る。",
    )
    return doc


def save(doc: Document, name: str) -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / name
    doc.save(str(path))
    print("wrote", path)
    return path


def main() -> None:
    save(build_ui_manual(), "データ集約_シナリオ編集画面説明書_V11106.docx")
    save(build_keys_concept(), "データ集約_主キー・連携キー・結合キーの動作概念_V11106.docx")
    save(build_json_doc(), "データ集約_シナリオファイルJson構成と概要_V11106.docx")
    save(build_dsl_doc(), "データ集約_整形DSL_コマンドリファレンス_V11106.docx")


if __name__ == "__main__":
    main()
